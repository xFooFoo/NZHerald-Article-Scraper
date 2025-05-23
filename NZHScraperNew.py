import requests
import re
import os
from docx import Document
from bs4 import BeautifulSoup



file_type = '.docx' #feel free to replace this with .doc, .docx, .txt etc...
url = input('What\'s the URL for the NZ Herald Article? ')
#url = 'https://www.nzherald.co.nz/business/we-didnt-have-any-food-in-the-house-asb-ceo-vittoria-shortt-on-her-tough-childhood-money-talks/TNURNHWFOBATTMFZF6EPR3ZSCM/#:~:text=ASB%20bank%20chief%20executive%20Vittoria,bank%20chief%20executive%20Vittoria%20Shortt.'
response = requests.get(url)
content = ""

def writeToTxtFile(file_path, content):
    with open(file_path, 'w', encoding='utf-8', newline='') as file:
        # Write the content to the same location as the script
        file.write(content)

def writeToDocxFile(file_path, content):
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.strip().split('\n\n'):
        doc.add_paragraph(paragraph)
    doc.save(file_path)

writeToFile = {
    ".docx": writeToDocxFile,
    ".txt": writeToTxtFile,
}

# Check if the request was successful
if response.status_code == 200:
    # Parse the HTML content of the page
    soup = BeautifulSoup(response.text, 'html.parser')

    # Title Portion
    title = str(soup.find(class_='article__heading').text)
    content += "Title: " + title + "\n\n"

    # Sometimes the article only has one body section
    parent_section = soup.select_one('section[data-test-ui="article__body"]')
    if parent_section:
        paragraphs = parent_section.find_all('p')
        # Extract and print the text content of each element
        for paragraph in paragraphs:
            content += paragraph.text + "\n\n"

    # This is usually the article preview part
    parent_section = soup.select_one('section[data-test-ui="article-top-body"]')
    if parent_section:
        paragraphs = parent_section.find_all('p')
        # Extract and print the text content of each element
        for paragraph in paragraphs:
            content += paragraph.text + "\n\n"

    # This is usually the latter part of the article where it's hidden behind a paywall
    parent_section = soup.select_one('section[data-test-ui="article-bottom-body"]')
    if parent_section:
        paragraphs = parent_section.find_all('p')
        # Extract and print the text content of each element
        for paragraph in paragraphs:
            content += paragraph.text + "\n\n"

    ############ Write content to a txt file ##########################
    # Get the absolute path of the current script
    script_directory = os.path.abspath(os.path.dirname(__file__))
    # filename can't have these special chars
    file_path = script_directory + "\\" + re.sub("[\\\\/:*?\"<>|]", "", title) + file_type
    # Get the writer function
    writer = writeToFile.get(file_type)
    if writer:
        writer(file_path, content)
        print(f"Content written to \"{file_path}\"")
    else:
        print(f"Failed to write: unsupported file type '{file_type}'")
else:
    print(f"Failed to retrieve the page. Status code: {response.status_code}")

#print(content)