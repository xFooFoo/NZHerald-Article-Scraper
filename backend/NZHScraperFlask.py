import requests
import re
import os
from docx import Document
from bs4 import BeautifulSoup

def scrapeContent(url):
    response = requests.get(url)
    content = []
    # file_type = '.docx' #feel free to replace this with .doc, .docx, .txt etc...

    # def writeToTxtFile(file_path, content):
    #     with open(file_path, 'w', encoding='utf-8', newline='') as file:
    #         # Write the content to the same location as the script
    #         file.write(content)

    # def writeToDocxFile(file_path, content):
    #     doc = Document()
    #     doc.add_heading(title, level=1)
    #     for paragraph in content.strip().split('\n\n'):
    #         doc.add_paragraph(paragraph)
    #     doc.save(file_path)

    # writeToFile = {
    #     ".docx": writeToDocxFile,
    #     ".txt": writeToTxtFile,
    # }

    # Check if the request was successful
    if response.status_code == 200:
        # Parse the HTML content of the page
        soup = BeautifulSoup(response.text, 'html.parser')

        # Title Portion
        heading = soup.select_one('h1[data-test-ui="article__heading"]')
        if heading:
            title = heading.text 
        else:
            title = "Title not found"
        # content += "Title: " + title + "\n\n"

        # Sometimes the article only has one body section
        parent_section = soup.select_one('section[data-test-ui="article__body"]')
        if parent_section:
            paragraphs = parent_section.find_all('p')
            # Extract and print the text content of each element
            for paragraph in paragraphs:
                content.append(paragraph.text)

        # This is usually the article preview part
        parent_section = soup.select_one('section[data-test-ui="article-top-body"]')
        if parent_section:
            paragraphs = parent_section.find_all('p')
            # Extract and print the text content of each element
            for paragraph in paragraphs:
                content.append(paragraph.text)

        # This is usually the latter part of the article where it's hidden behind a paywall
        parent_section = soup.select_one('section[data-test-ui="article-bottom-body"]')
        if parent_section:
            paragraphs = parent_section.find_all('p')
            # Extract and print the text content of each element
            for paragraph in paragraphs:
                content.append(paragraph.text)

        ############ Write content to a txt file ##########################
        # Get the absolute path of the current script
        # script_directory = os.path.abspath(os.path.dirname(__file__))
        # filename can't have these special chars
        # file_path = script_directory + "\\" + re.sub("[\\\\/:*?\"<>|]", "", title) + file_type
        # Get the writer function
        # writer = writeToFile.get(file_type)
        # if writer:
        #     writer(file_path, content)
        #     print(f"Content written to \"{file_path}\"")
        # else:
        #     print(f"Failed to write: unsupported file type '{file_type}'")
    else:
        print(f"Failed to retrieve the page. Status code: {response.status_code}")
    return title, content
    #print(content)